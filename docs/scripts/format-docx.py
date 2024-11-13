import os
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Function to add a footer with the current date
def add_footer_with_date(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    
    # Add text and apply formatting
    current_date = datetime.now().strftime("%Y-%m-%d")
    run = paragraph.add_run(f"Generated on {current_date}")
    
    # Apply font size, color, and italic style
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True

# Define folder path and list of search strings
folder_path = "docx"
search_strings = [
    "      sed -i \"s/^\\(CONFIG_MODULE_SIG_FORCE\\).*/\\\\1 policy\\<{'arm64': 'n', 'armhf': 'n'}\\>/\" ${DEBIAN}/config/annotations",
    "  sed -i \"s/^\\(CONFIG_MODULE_SIG_ALL.*\\)'arm64': 'y'\\(.*\\)/\\\\1'arm64': 'n'\\\\2/\" ${DEBIAN}/config/annotations",
    "sudo snapcraft --target-arch=arm64 --destructive-mode --enable-experimental-target-arch",
    "sudo dpkg -i linux-image-unsigned-<kernel version>-<generic or derivative>*.deb",
    "sudo dpkg -i linux-headers-<kernel version>-<generic or derivative>*.deb",
    "sudo dpkg -i linux-modules-<kernel version>-<generic or derivative>*.deb",
    "deb [arch=amd64] http://archive.ubuntu.com/ubuntu focal main restricted",
    "<kernel_source_working_directory>/debian.<derivative>/changelog",
    "snap install --dangerous --devmode <name>_<version>_<arch>.snap",
    "linux-image-unsigned-6.8.0-999-generic_6.8.0-999.48_amd64.deb",
    "sudo apt build-dep -y linux linux-image-unsigned-$(uname -r)",
    "deb-src http://archive.ubuntu.com/ubuntu jammy-updates main",
    "Signed-By: /usr/share/keyrings/ubuntu-archive-keyring.gpg",
    "<kernel_source_working_directory>/debian.master/changelog",
    "sudo apt install -y fakeroot llvm libncurses-dev dwarves",
    "linux-headers-6.8.0-999-generic_6.8.0-999.48_amd64.deb",
    "linux-modules-6.8.0-999-generic_6.8.0-999.48_amd64.deb",
    "sudo dpkg -i linux-headers-<kernel version>*_all.deb",
    "deb-src http://archive.ubuntu.com/ubuntu jammy main",
    "sudo snapcraft --build-for=arm64 --destructive-mode",
    "Components: main universe restricted multiverse",
    "ln -s snap/local/<project>.yaml snapcraft.yaml",
    "linux-headers-6.8.0-999_6.8.0-999.48_all.deb",
    "Suites: noble noble-updates noble-backports",
    "apt source linux-image-unsigned-$(uname -r)",
    "linux (6.8.0-999.48) noble; urgency=medium",
    "/etc/apt/sources.list.d/ubuntu.sources",
    "URIs: http://archive.ubuntu.com/ubuntu",
    "sudo snap install snapcraft --classic",
    "cd <kernel_source_working_directory>",
    "git clone <kernel-source-repository>",
    "sudo dpkg --add-architecture arm64",
    "dpkg --print-foreign-architectures",
    "fakeroot debian/rules editconfigs",
    "chmod a+x debian/scripts/misc/*",
    "User <your Launchpad username>",
    "cd <kernel-source-repository>",
    "<name>_<version>_<arch>.snap",
    "fakeroot debian/rules binary",
    "└── linux_X.Y.Z.orig.tar.gz",
    "fakeroot debian/rules clean",
    "sudo apt purge -y snapcraft",
    "chmod a+x debian/scripts/*",
    "├── linux_X.Y.Z-*.diff.gz",
    "      # override configs",
    "sudo apt-get -y upgrade",
    "chmod a+x debian/rules",
    "Types: deb deb-src",    
    "Host git.launchpad.net",
    "CONFIG_MODULE_SIG_ALL",
    "├── linux_X.Y.Z-*.dsc",
    "/etc/apt/sources.list",
    "    override-build: |",
    "Architectures: amd64",
    "<working_directory>",
    "sudo apt-get update",
    "git.launchpad.net",
    "├── linux-X.Y.Z/",
    "sudo apt update",
    "ubuntu.sources",
    "snapcraft.yaml",
    "~/.ssh/config",
    "'arm64': 'n'",
    "sources.list",
    "sudo reboot",
    "menuconfig",
    "snap/local",
    "│   └── *",
    "  kernel:",
    "uname -r",
    "dpkg -i",
    "Types:",
    "initrd",
    "parts:",
    "[...]",
]

# Step 1: List all .docx files in the folder
print("step 1")
docx_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]

for filename in docx_files:
    file_path = os.path.join(folder_path, filename)

    # Open the original .docx file for in-place modification
    doc = Document(file_path)

    # Process each paragraph in the document
    for para in doc.paragraphs:
        for run in para.runs:
            # Check if any of the search strings are in the current run's text
            if any(search_string in run.text for search_string in search_strings):
                # Format matching text
                run.font.name = 'Ubuntu Mono'
                run.font.size = Pt(11)

                # Apply background color #EEEDEB
                rPr = run._r.get_or_add_rPr()  # Access the run properties
                shading_elm = parse_xml(r'<w:shd {} w:fill="EEEDEB"/>'.format(nsdecls('w')))
                rPr.append(shading_elm)

    # Call the footer-adding function
    add_footer_with_date(doc)

    # Save changes directly to the original file
    doc.save(file_path)
    print(f"Formatted file: {file_path}")

# Move the source folder to the destination folder
shutil.move(folder_path, os.path.join("_build", os.path.basename(folder_path)))
